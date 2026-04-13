# -*- coding: utf-8 -*-
"""
论文格式助手 v1.0
一款通用的论文格式自动排版工具
"""

import re
import os
import subprocess
import traceback
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTableWidget, QHeaderView, QPushButton,
    QLabel, QLineEdit, QProgressBar, QFileDialog, QMessageBox,
    QFrame, QCheckBox, QComboBox, QStyleFactory
)
from PySide6.QtCore import Qt, QThread, Signal

# ==================== 配置数据 ====================

# 中文字体选项
CHINESE_FONTS = [
    "宋体", "黑体", "楷体", "仿宋", "微软雅黑", "华文宋体", "华文黑体",
    "华文楷体", "华文仿宋", "方正宋体", "方正黑体"
]

# 英文字体选项
ENGLISH_FONTS = [
    "Times New Roman", "Arial", "Calibri", "Georgia",
    "Palatino Linotype", "Book Antiqua"
]

# 字号选项（名称 -> 磅值）
FONT_SIZES = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
    "六号": 7.5, "七号": 5.5, "八号": 5
}
FONT_SIZE_NAMES = list(FONT_SIZES.keys())

# 对齐方式
ALIGNMENTS = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY
}
ALIGNMENT_NAMES = list(ALIGNMENTS.keys())

# 行距选项（倍数和固定值）
LINE_SPACINGS = {
    "单倍行距": 1.0,
    "1.15倍行距": 1.15,
    "1.25倍行距": 1.25,
    "1.5倍行距": 1.5,
    "1.75倍行距": 1.75,
    "2倍行距": 2.0,
    "2.5倍行距": 2.5,
    "3倍行距": 3.0,
    "6磅": 6,
    "12磅": 12,
    "15磅": 15,
    "18磅": 18,
    "20磅": 20,
    "21磅": 21,
    "22磅": 22,
    "24磅": 24,
    "28磅": 28,
    "30磅": 30
}
LINE_SPACING_NAMES = list(LINE_SPACINGS.keys())

# 段前段后间距选项（磅）
SPACINGS = ["0磅", "6磅", "12磅", "18磅", "24磅", "30磅", "36磅", "48磅"]

# 首行缩进选项
FIRST_LINE_INDENTS = ["无", "1字符", "2字符", "3字符", "4字符"]

# 整体缩进选项
OVERALL_INDENTS = [
    "无",
    "左缩进1字符", "左缩进2字符", "左缩进3字符", "左缩进4字符",
    "右缩进1字符", "右缩进2字符",
    "左右各1字符", "左右各2字符"
]

DEFAULT_OUTPUT_DIR_HINT = "与原文件相同目录"
DEFAULT_FILENAME_HINT = "留空则使用默认名称"

# 论文部分列表
PARTS = [
    ("摘要标题", "abstract_title"),
    ("摘要内容", "abstract_content"),
    ("一级标题", "heading1"),
    ("二级标题", "heading2"),
    ("三级标题", "heading3"),
    ("四级标题", "heading4"),
    ("正文内容", "body"),
    ("图片标题", "figure_caption"),
    ("表格标题", "table_caption"),
    ("表格内容", "table_content"),
    ("图表注释", "table_note"),
    ("公式", "formula"),
    ("参考文献标题", "ref_title"),
    ("参考文献内容", "ref_content"),
    ("致谢标题", "ack_title"),
    ("致谢内容", "ack_content"),
    ("附录标题", "appendix_title"),
    ("附录内容", "appendix_content"),
]

TABLE_COLUMNS = [
    ("part_name", "", 150),
    ("chinese_font", "中文字体", 130),
    ("english_font", "英文字体", 150),
    ("font_size", "字号", 88),
    ("alignment", "对齐", 88),
    ("line_spacing", "行距", 118),
    ("space_before", "段前", 88),
    ("space_after", "段后", 88),
    ("first_line_indent", "首行缩进", 104),
    ("overall_indent", "整体缩进", 124),
    ("bold", "加粗", 78),
    ("italic", "倾斜", 78),
]

# ==================== 识别规则 ====================

PART_STYLE_NAMES = {part_key: part_name for part_name, part_key in PARTS}

OUTLINE_LEVELS = {
    "abstract_title": 0,
    "heading1": 0,
    "heading2": 1,
    "heading3": 2,
    "heading4": 3,
    "ref_title": 0,
    "appendix_title": 0,
}

FIRST_LINE_INDENT_MAP = {
    "无": 0,
    "1字符": 1,
    "2字符": 2,
    "3字符": 3,
    "4字符": 4,
}

OVERALL_INDENT_MAP = {
    "无": (0, 0),
    "左缩进1字符": (1, 0),
    "左缩进2字符": (2, 0),
    "左缩进3字符": (3, 0),
    "左缩进4字符": (4, 0),
    "右缩进1字符": (0, 1),
    "右缩进2字符": (0, 2),
    "左右各1字符": (1, 1),
    "左右各2字符": (2, 2),
}


class PartIdentifier:
    """论文部分识别器"""

    # 一级标题：第X章（支持中文数字和阿拉伯数字）
    HEADING1_PATTERN = re.compile(r'^第[一二三四五六七八九十百0-9]+章\s*.*$')

    # 二级标题：第X节（支持中文数字和阿拉伯数字）
    HEADING2_PATTERN = re.compile(r'^第[一二三四五六七八九十百0-9]+节\s*.*$')

    # 三级标题：
    # - 1.1.1 ××× 格式（阿拉伯数字，三个数字用.分隔）
    # - 一、××× 格式（中文数字后跟、）
    HEADING3_PATTERN = re.compile(r'^(?:[0-9]+\.[0-9]+\.[0-9]+\s+|[一二三四五六七八九十百]+、)\s*.*$')

    # 四级标题：
    # - 1.1.1.1 ××× 格式（阿拉伯数字，四个数字用.分隔）
    # - （一）××× 格式（中文数字在括号内）
    HEADING4_PATTERN = re.compile(r'^(?:[0-9]+\.[0-9]+\.[0-9]+\.[0-9]+\s+|（[一二三四五六七八九十百]+）)\s*.*$')

    # 阿拉伯数字编号的标题（如1. 2. 1.1 2.1等）
    ARABIC_HEADING1_PATTERN = re.compile(r'^[0-9]+\s+.*$')  # 1 人民币国际化
    ARABIC_HEADING1_DOT_PATTERN = re.compile(r'^[0-9]+\s*[、．．]\s*.*$')  # 1、或1．
    ARABIC_HEADING2_PATTERN = re.compile(r'^[0-9]+\.[0-9]+\s*.*$')  # 1.1
    ARABIC_HEADING3_PATTERN = re.compile(r'^[0-9]+\.[0-9]+\.[0-9]+\s*.*$')  # 1.1.1

    # 图片标题：图X-X、图X.X、图X等（宽松的匹配）
    FIGURE_PATTERN = re.compile(r'^图\s*[0-9]+([-\.][0-9]+)?.*$')

    # 表格标题：表X-X、表X.X、表X等（宽松的匹配）
    TABLE_PATTERN = re.compile(r'^表\s*[0-9]+([-\.][0-9]+)?.*$')

    # 表格注释：注：或 注:
    TABLE_NOTE_PATTERN = re.compile(r'^注[：:].*$')
    FORMULA_NUMBER_PATTERN = re.compile(r'^\s*[\(\uff08]?\s*\d+(?:[-\.]\d+)*\s*[\)\uff09]?\s*$')

    # 特殊标题（允许中间有空格）
    @staticmethod
    def _normalize(text):
        """去除所有空格用于匹配"""
        return text.replace(" ", "").replace("\u3000", "")

    @staticmethod
    def _has_formula_omml(paragraph):
        """检测段落中是否包含 Word 插入公式生成的 OMML 节点"""
        try:
            return bool(paragraph._element.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]'))
        except Exception:
            return False

    @staticmethod
    def _is_formula_paragraph(paragraph, text):
        """判断是否为独立公式段落"""
        if not PartIdentifier._has_formula_omml(paragraph):
            return False

        normalized = PartIdentifier._normalize(text)

        # Word 公式在 python-docx 中经常显示为空文本
        if not normalized:
            return True

        # 仅带公式编号的段落
        if PartIdentifier.FORMULA_NUMBER_PATTERN.match(normalized):
            return True

        # 兼容"公式 + 编号"这类短文本段落，避免把正文中的行内公式误判成整段公式
        non_formula_text = re.sub(r'[\s\(\)\uff08\uff090-9\-\.=+*/<>≤≥,，:：;；]+', '', normalized)
        return len(non_formula_text) <= 6

    @staticmethod
    def identify(paragraph, position_context=None):
        """
        识别段落类型
        返回: (part_type, is_title)
        """
        text = paragraph.text.strip()

        if PartIdentifier._is_formula_paragraph(paragraph, text):
            return "formula", False

        # 空段落
        if not text:
            return None, False

        normalized = PartIdentifier._normalize(text)

        # 改进的单独一行判断：更智能的标题识别逻辑
        has_end_punctuation = text.endswith(('。', '！', '？', '.', '!', '?', '；', ';', '：', ':'))
        # 更宽松的长度限制，因为中文标题可能较长
        is_short_text = len(text) < 150
        # 检查是否包含换行符或其他段落特征
        has_line_breaks = '\n' in text or '\r' in text
        # 检查是否像标题（包含标题关键词且不以句号结束）
        contains_title_keywords = any(keyword in text for keyword in ['第', '章', '节', '、', '（', '）', '图', '表'])

        # 标题字数限制：超过50个字符（不含空格）不算标题
        char_count_without_spaces = len(text.replace(" ", "").replace("\u3000", ""))
        is_too_long_for_title = char_count_without_spaces > 50

        is_single_line = (is_short_text and not has_end_punctuation and not has_line_breaks and not is_too_long_for_title) or \
                        (contains_title_keywords and not is_too_long_for_title)

        # 1. 特殊标题检查（支持更多变体）
        if normalized in ["摘要", "摘 要", "ABSTRACT", "Abstract"]:
            return "abstract_title", True
        if normalized in ["参考文献", "参 考 文 献", "REFERENCE", "Reference", "REFERENCES", "References", "οο"]:
            return "ref_title", True
        if normalized in ["致谢", "致 谢", "谢辞", "ACKNOWLEDGEMENT", "Acknowledgement", "ACKNOWLEDGEMENTS", "Acknowledgements"]:
            return "ack_title", True
        if normalized in ["附录", "附 录", "APPENDIX", "Appendix", "APPENDICES", "Appendices"]:
            return "appendix_title", True

        # 2. 标题级别检查（需要单独一行）
        if is_single_line:
            if PartIdentifier.HEADING1_PATTERN.match(text):
                return "heading1", True
            if PartIdentifier.HEADING2_PATTERN.match(text):
                return "heading2", True
            if PartIdentifier.HEADING3_PATTERN.match(text):
                return "heading3", True
            if PartIdentifier.HEADING4_PATTERN.match(text):
                return "heading4", True

            # 阿拉伯数字编号的标题
            if PartIdentifier.ARABIC_HEADING2_PATTERN.match(text) or PartIdentifier.ARABIC_HEADING3_PATTERN.match(text):
                return "heading2", True
            if PartIdentifier.ARABIC_HEADING1_PATTERN.match(text) or PartIdentifier.ARABIC_HEADING1_DOT_PATTERN.match(text):
                return "heading1", True

        # 3. 图表标题检查（需要单独一行）
        if is_single_line:
            if PartIdentifier.FIGURE_PATTERN.match(text):
                return "figure_caption", True
            if PartIdentifier.TABLE_PATTERN.match(text):
                return "table_caption", True

        # 4. 表格注释
        if PartIdentifier.TABLE_NOTE_PATTERN.match(text):
            return "table_note", True

        # 5. 根据上下文和内容判断区域
        if position_context:
            if position_context == "abstract":
                return "abstract_content", False
            elif position_context == "ref":
                return "ref_content", False
            elif position_context == "ack":
                return "ack_content", False
            elif position_context == "appendix":
                return "appendix_content", False

        # 5. 检查是否为参考文献内容（在参考文献标题后且不是其他标题）
        if position_context == "ref":
            return "ref_content", False

        # 6. 默认为正文
        return "body", False


# ==================== 样式管理器 ====================

class StyleManager:
    """样式管理器"""

    def __init__(self, document):
        self.document = document
        self.styles = document.styles

    def get_style_name(self, part_type):
        return PART_STYLE_NAMES.get(part_type, part_type)

    def _set_style_alignment(self, style, alignment_name):
        alignment_value = ALIGNMENTS.get(alignment_name)
        if alignment_value is None:
            return

        style.paragraph_format.alignment = alignment_value

        if style.paragraph_format.alignment != alignment_value:
            style.element.get_or_add_pPr().jc_val = alignment_value

    def _set_style_outline_level(self, style, part_type):
        ppr = style.element.get_or_add_pPr()
        outline_level = OUTLINE_LEVELS.get(part_type)

        if outline_level is None:
            if ppr.outlineLvl is not None:
                ppr._remove_outlineLvl()
            return

        ppr.get_or_add_outlineLvl().val = outline_level

    def create_or_update_style(self, part_type, rule):
        """
        创建或更新样式
        part_type: 段落类型（如"heading1"）
        rule: 格式规则字典
        """
        style_name = self.get_style_name(part_type)

        # 检查样式是否存在
        if style_name in self.styles:
            style = self.styles[style_name]
            print(f"更新样式: {style_name}")
        else:
            # 创建新样式（基于Normal样式）
            style = self.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            print(f"创建样式: {style_name}")

        # 设置段落格式
        self._set_paragraph_format(style, rule)
        self._set_style_outline_level(style, part_type)

        # 设置字体格式
        self._set_font_format(style, rule)

        return style

    def _set_paragraph_format(self, style, rule):
        """设置段落格式"""
        pf = style.paragraph_format

        # 对齐方式
        if rule.get("alignment"):
            self._set_style_alignment(style, rule["alignment"])

        # 行距
        if rule.get("line_spacing"):
            line_spacing_map = {
                "单倍行距": 1.0, "1.15倍行距": 1.15, "1.25倍行距": 1.25,
                "1.5倍行距": 1.5, "1.75倍行距": 1.75, "2倍行距": 2.0,
                "2.5倍行距": 2.5, "3倍行距": 3.0
            }
            if rule["line_spacing"] in line_spacing_map:
                pf.line_spacing = line_spacing_map[rule["line_spacing"]]

        # 段前间距
        if rule.get("space_before"):
            try:
                space_value = int(rule["space_before"].rstrip('磅'))
                pf.space_before = Pt(space_value)
            except:
                pass

        # 段后间距
        if rule.get("space_after"):
            try:
                space_value = int(rule["space_after"].rstrip('磅'))
                pf.space_after = Pt(space_value)
            except:
                pass

        # 整体缩进
        overall_indent = rule.get("overall_indent", "无")
        left_chars, right_chars = OVERALL_INDENT_MAP.get(overall_indent, (0, 0))
        pf.left_indent = Pt(left_chars * 12)
        pf.right_indent = Pt(right_chars * 12)

        # 首行缩进
        first_line_indent = rule.get("first_line_indent", "无")
        indent_chars = FIRST_LINE_INDENT_MAP.get(first_line_indent, 0)
        pf.first_line_indent = Pt(indent_chars * 12)  # 约12磅=1字符

    def _set_font_format(self, style, rule):
        """设置字体格式"""
        font = style.font

        # 中文字体
        if rule.get("chinese_font"):
            font.name = rule["chinese_font"]
            # 设置西文字体
            font._element.rPr.rFonts.set(qn('w:eastAsia'), rule["chinese_font"])

        # 英文字体
        if rule.get("english_font"):
            font.name = rule["english_font"]

        # 字号
        if rule.get("font_size"):
            size_map = {
                "初号": 42, "小初": 36, "一号": 26, "小一": 24,
                "二号": 22, "小二": 18, "三号": 16, "小三": 15,
                "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
                "六号": 7.5, "七号": 5.5, "八号": 5
            }
            if rule["font_size"] in size_map:
                font.size = Pt(size_map[rule["font_size"]])

        # 加粗
        if rule.get("bold") is not None:
            font.bold = rule["bold"]

        # 倾斜
        if rule.get("italic") is not None:
            font.italic = rule["italic"]

    def apply_style_to_paragraph(self, paragraph, part_type, rules):
        """
        将样式应用到段落
        策略：清除原有直接格式，应用样式
        """
        if part_type is None:
            # 如果part_type是None，默认为正文内容
            part_type = "body"

        style_name = self.get_style_name(part_type)

        # 总是更新样式以确保与当前GUI规则一致
        # 这样即使样式已存在，也会按最新的规则刷新其格式
        target_style = self.create_or_update_style(part_type, rules.get(part_type, {}))

        # 应用样式到段落
        paragraph.style = target_style

        # 清除原有的直接格式（可选策略）
        # 这里我们选择完全依赖样式，所以清除直接格式
        try:
            paragraph.paragraph_format.clear()
        except:
            pass

        # 确保对齐方式正确设置（因为clear()可能不会正确继承样式对齐）
        style_pf = target_style.paragraph_format
        if hasattr(style_pf, 'alignment') and style_pf.alignment is not None:
            paragraph.paragraph_format.alignment = style_pf.alignment

        # 谨慎清除run上的直接格式，避免破坏图片和公式
        for run in paragraph.runs:
            try:
                # 检查是否是图片或公式（这些节点不应被当作纯文本run清理）
                has_drawing = False
                has_formula = False
                if hasattr(run, '_element') and run._element is not None:
                    try:
                        drawings = run._element.xpath('.//w:drawing')
                        has_drawing = len(drawings) > 0
                        formulas = run._element.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]')
                        has_formula = len(formulas) > 0
                    except:
                        pass

                if has_drawing or has_formula:
                    # 对于包含图片或公式的run，只清除安全的字符级属性
                    run.font.bold = None
                    run.font.italic = None
                    run.font.underline = None
                    run.font.color = None
                else:
                    # 对于纯文本run，清除所有格式
                    run.font.name = None
                    run.font.size = None
                    run.font.bold = None
                    run.font.italic = None
                    run.font.underline = None
                    run.font.color = None
            except:
                pass

    def apply_style_to_table_cell(self, cell, part_type, rules):
        """将样式应用到表格单元格中的段落"""
        for paragraph in cell.paragraphs:
            self.apply_style_to_paragraph(paragraph, part_type, rules)


# ==================== 格式应用器（新版） ====================

class FormatApplier:
    """格式应用器（新版 - 样式驱动）"""

    def __init__(self, format_rules, document=None):
        self.rules = format_rules
        self.document = document
        self.style_manager = StyleManager(document) if document else None

    def apply_to_paragraph(self, paragraph, part_type):
        """对段落应用样式"""
        # 使用样式管理器
        if self.style_manager:
            self.style_manager.apply_style_to_paragraph(paragraph, part_type, self.rules)
        else:
            # 回退到直接格式（如果样式管理器不可用）
            self._apply_direct_format(paragraph, part_type)

    def _apply_direct_format(self, paragraph, part_type):
        """直接格式应用（回退方法）"""
        # 确保至少应用正文格式，即使part_type不在规则中
        if part_type not in self.rules:
            if "body" in self.rules:
                rule = self.rules["body"]
            else:
                return
        else:
            rule = self.rules[part_type]

        if not rule:
            return

        # 获取段落格式
        pf = paragraph.paragraph_format

        # 设置对齐方式
        if rule.get("alignment"):
            pf.alignment = ALIGNMENTS.get(rule["alignment"])

        # 设置行距
        if rule.get("line_spacing"):
            line_spacing_value = LINE_SPACINGS.get(rule["line_spacing"])
            if line_spacing_value:
                # 检查是否为固定值（包含"="磅"）
                if "倍" in rule["line_spacing"]:
                    # 倍数行距
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = line_spacing_value
                else:
                    # 固定值行距
                    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    pf.line_spacing = Pt(line_spacing_value)

        # 设置段前间距
        if rule.get("space_before"):
            space_value = self._parse_spacing(rule["space_before"])
            if space_value is not None:
                pf.space_before = Pt(space_value)

        # 设置段后间距
        if rule.get("space_after"):
            space_value = self._parse_spacing(rule["space_after"])
            if space_value is not None:
                pf.space_after = Pt(space_value)

        # 设置整体缩进
        left_chars, right_chars = self._parse_overall_indent(rule.get("overall_indent"))
        pf.left_indent = Pt(left_chars * 12)
        pf.right_indent = Pt(right_chars * 12)

        # 设置首行缩进
        indent_chars = self._parse_first_line_indent(rule.get("first_line_indent"))
        pf.first_line_indent = Pt(indent_chars * 12)

        # 设置字体格式
        for run in paragraph.runs:
            self._apply_font(run, rule)

    def _apply_font(self, run, rule):
        """设置字体格式"""
        # 中文字体
        if rule.get("chinese_font"):
            run.font.name = rule["chinese_font"]
            run._element.rPr.rFonts.set(qn('w:eastAsia'), rule["chinese_font"])

        # 英文字体
        if rule.get("english_font"):
            run.font.name = rule["english_font"]

        # 字号
        if rule.get("font_size"):
            size_pt = FONT_SIZES.get(rule["font_size"])
            if size_pt:
                run.font.size = Pt(size_pt)

        # 加粗
        if rule.get("bold") is not None:
            run.font.bold = rule["bold"]

        # 倾斜
        if rule.get("italic") is not None:
            run.font.italic = rule["italic"]

    def _parse_spacing(self, spacing_str):
        """解析间距字符串，返回磅值"""
        if not spacing_str:
            return None
        match = re.match(r'(\d+)', spacing_str)
        if match:
            return int(match.group(1))
        return None

    def _parse_first_line_indent(self, indent_str):
        """解析首行缩进字符串，返回字符数"""
        if not indent_str:
            return 0
        return FIRST_LINE_INDENT_MAP.get(indent_str, 0)

    def _parse_overall_indent(self, indent_str):
        """解析整体缩进字符串，返回左右缩进字符数"""
        if not indent_str:
            return (0, 0)
        return OVERALL_INDENT_MAP.get(indent_str, (0, 0))

    def apply_to_table(self, table, rules):
        """对表格应用样式（增强版）"""
        if not self.style_manager:
            return

        # 获取PartIdentifier用于识别（避免循环导入）
        identifier = PartIdentifier()

        # 遍历表格的每一行和单元格
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_paragraphs = list(cell.paragraphs)  # 复制列表，避免修改迭代

                # 对单元格中的每个段落进行识别和应用
                for para_idx, paragraph in enumerate(cell_paragraphs):
                    # 重新获取文本，因为段落对象可能在处理过程中被修改
                    text = paragraph.text.strip()

                    # 跳过空段落
                    if not text:
                        continue

                    # 识别优先级（从高到低）：
                    # 1. 图片标题（如果单元格中有图片且匹配图标题模式）
                    # 2. 表格标题
                    # 3. 图表注释（资料来源、数据来源等）
                    # 4. 其他（表格内容）

                    # 检查是否为表格标题或图片标题
                    chart_result = self._is_chart_caption(text)
                    if chart_result and chart_result[0]:
                        caption_type = chart_result[1]
                        self.style_manager.apply_style_to_paragraph(
                            paragraph, caption_type, rules
                        )
                        continue

                    # 检查是否为图表注释（资料来源、数据来源等）
                    if (text.startswith("资料来源") or
                        text.startswith("数据来源") or
                        text.startswith("来源") or
                        text.startswith("注")):

                        self.style_manager.apply_style_to_paragraph(
                            paragraph, "table_note", rules
                        )
                        continue

                    # 使用段落识别器判断其他类型
                    part_type, is_title = identifier.identify(paragraph)

                    # 应用对应的样式
                    if part_type and part_type in ["figure_caption", "table_caption", "table_note"]:
                        self.style_manager.apply_style_to_paragraph(paragraph, part_type, rules)
                    else:
                        # 默认为表格内容
                        self.style_manager.apply_style_to_paragraph(paragraph, "table_content", rules)

    def _is_chart_caption(self, text):
        """检查是否为图表标题"""
        # 表标题模式（支持各种变体）
        table_patterns = [
            r'^表\s*[0-9]+[-\.][0-9]+.*$',      # 表5.1...
            r'^表\s*[0-9]+[-\.][0-9]+\s*[:：].*$', # 表5.1: ...
            r'^表\s*[0-9]+.*$',                 # 表5...
            r'^表[0-9]+[-\.][0-9]+.*$',         # 表5.1
            r'^表[0-9]+.*$',                    # 表5
            r'^表\s+[0-9]+.*$',                # 表 5.1
        ]

        # 图标题模式（支持各种变体）
        figure_patterns = [
            r'^图\s*[0-9]+[-\.][0-9]+.*$',      # 图1.1...
            r'^图\s*[0-9]+[-\.][0-9]+\s*[:：].*$', # 图1.1: ...
            r'^图\s*[0-9]+.*$',                 # 图1...
            r'^图[0-9]+[-\.][0-9]+.*$',         # 图1.1
            r'^图[0-9]+.*$',                    # 图1
            r'^图\s+[0-9]+.*$',                # 图 1.1
        ]

        # 图表注释模式（单独处理，不与图表标题冲突）
        annotation_patterns = [
            r'^资料来源.*$',                   # 资料来源
            r'^数据来.*$',                     # 数据来源
            r'^来源.*$',                       # 来源
            r'^注.*$',                         # 注
        ]

        import re
        for pattern in table_patterns:
            if re.match(pattern, text):
                return (True, "table_caption")

        for pattern in figure_patterns:
            if re.match(pattern, text):
                return (True, "figure_caption")

        # 检查是否为图表注释（单独处理）
        for pattern in annotation_patterns:
            if re.match(pattern, text):
                return (True, "table_note")

        return None


# ==================== 自定义组合框和复选框 widget ====================

class TableComboBox(QWidget):
    """表格中使用的组合框"""
    def __init__(self, items, default_item="", parent=None):
        super().__init__(parent)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(4, 4, 4, 4)

        self.combo = QComboBox()
        self.combo.addItems(items)
        if default_item in items:
            self.combo.setCurrentText(default_item)
            
        # 匹配图中的精致下拉框：纯白背景，浅灰边框，圆角，字号适中
        self.combo.setStyleSheet("""
            QComboBox {
                background-color: #FFFFFF;
                border: 1px solid #CBD5E1;
                border-radius: 6px;
                padding: 4px 4px;
                font-family: "Microsoft YaHei UI";
                font-size: 13px;
                color: #334155;
            }
            QComboBox:hover {
                border: 1px solid #93C5FD;
                background-color: #F8FAFC;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 0px;
                border-left: none;
            }
            QComboBox::down-arrow {
                width: 0px;
                height: 0px;
            }
        """)

        self.combo.installEventFilter(self)
        layout.addWidget(self.combo)

    def text(self): return self.combo.currentText()
    def setText(self, text):
        index = self.combo.findText(text)
        if index >= 0: self.combo.setCurrentIndex(index)
    def eventFilter(self, obj, event):
        from PySide6.QtCore import QEvent
        if event.type() == QEvent.Wheel: return True
        return super().eventFilter(obj, event)


class TableCheckBox(QWidget):
    """表格中使用的复选框"""
    def __init__(self, checked=False, parent=None):
        super().__init__(parent)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setAlignment(Qt.AlignCenter)

        self.checkbox = QCheckBox()
        self.checkbox.setChecked(checked)
        
        self.checkbox.setStyleSheet("""
            QCheckBox { spacing: 0px; }
            QCheckBox::indicator {
                width: 16px; height: 16px;
                border: 1px solid #CBD5E1; border-radius: 4px; background-color: #FFFFFF;
            }
            QCheckBox::indicator:checked { background-color: #3B82F6; border: 1px solid #3B82F6; }
        """)
        layout.addWidget(self.checkbox)

    def isChecked(self): return self.checkbox.isChecked()
    def setChecked(self, checked): self.checkbox.setChecked(checked)


# ==================== 主程序界面 ====================

class FormatThread(QThread):
    """格式化工作线程"""
    progress_update = Signal(int, str)
    finished = Signal(bool, str, str)
    error = Signal(str)

    def __init__(self, task_snapshot):
        super().__init__()
        self.task_snapshot = task_snapshot

    def run(self):
        error_entries = []
        report_path = self._build_error_report_path(self.task_snapshot)
        try:
            input_path = self.task_snapshot["input_path"]
            output_path = self.task_snapshot["output_path"]
            rules = self.task_snapshot["rules"]
            output_dir = self.task_snapshot["output_dir"]
            custom_filename = self.task_snapshot["custom_filename"]

            self.progress_update.emit(10, "正在读取文档...")

            # 读取文档
            doc = Document(input_path)

            # 获取格式规则
            applier = FormatApplier(rules, doc)
            identifier = PartIdentifier()

            # 统计信息
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)

            self.progress_update.emit(20, "正在识别文档结构...")

            # 处理段落
            current_context = None

            processed_count = 0
            error_count = 0

            for i, paragraph in enumerate(doc.paragraphs):
                # 识别段落类型
                part_type, is_title = identifier.identify(paragraph, current_context)

                # 更新上下文
                if is_title:
                    if part_type == "abstract_title":
                        current_context = "abstract"
                    elif part_type == "ref_title":
                        current_context = "ref"
                    elif part_type == "ack_title":
                        current_context = "ack"
                    elif part_type == "appendix_title":
                        current_context = "appendix"
                    elif part_type in ["heading1", "heading2", "heading3", "heading4"]:
                        current_context = "body"

                # 应用格式
                if part_type:
                    try:
                        applier.apply_to_paragraph(paragraph, part_type)
                    except Exception as e:
                        error_count += 1
                        self._record_error(
                            error_entries,
                            "段落样式应用",
                            e,
                            index=i + 1,
                            part_type=part_type,
                            text=paragraph.text,
                        )
                else:
                    try:
                        applier.apply_to_paragraph(paragraph, "body")
                    except Exception as e:
                        error_count += 1
                        self._record_error(
                            error_entries,
                            "段落回退样式应用",
                            e,
                            index=i + 1,
                            part_type="body",
                            text=paragraph.text,
                        )

                processed_count += 1

                # 定期更新进度
                if processed_count % 100 == 0 or i == len(doc.paragraphs) - 1:
                    progress_val = 20 + int((processed_count / total_paragraphs) * 60)
                    self.progress_update.emit(progress_val, f"正在处理段落 {processed_count}/{total_paragraphs}...")

            self.progress_update.emit(80, "正在处理表格...")

            # 处理表格
            for i, table in enumerate(doc.tables):
                try:
                    applier.apply_to_table(table, rules)
                except Exception as e:
                    error_count += 1
                    table_preview = ""
                    try:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    table_preview = cell.text.strip()
                                    break
                            if table_preview:
                                break
                    except Exception:
                        table_preview = ""
                    self._record_error(
                        error_entries,
                        "表格处理",
                        e,
                        index=i + 1,
                        part_type="table",
                        text=table_preview,
                    )

                progress = 80 + int((i + 1) / max(total_tables, 1) * 10)
                self.progress_update.emit(progress, f"正在处理表格 {i + 1}/{total_tables}...")

            self.progress_update.emit(90, "正在保存文档...")

            # 确定输出路径
            input_path = self.task_snapshot["input_path"]
            output_dir = self.task_snapshot["output_dir"]
            custom_filename = self.task_snapshot["custom_filename"]

            if output_dir and output_dir != "与原文件相同目录":
                if custom_filename and custom_filename != "留空则使用默认名称":
                    if not custom_filename.endswith('.docx'):
                        custom_filename += '.docx'
                    output_path = os.path.join(output_dir, custom_filename)
                else:
                    output_path = os.path.join(output_dir, os.path.basename(input_path))
            else:
                name, ext = os.path.splitext(input_path)
                if custom_filename and custom_filename != "留空则使用默认名称":
                    if not custom_filename.endswith('.docx'):
                        custom_filename += '.docx'
                    output_path = os.path.join(os.path.dirname(input_path), custom_filename)
                else:
                    output_path = f"{name}_formatted{ext}"

            os.makedirs(os.path.dirname(output_path) or os.getcwd(), exist_ok=True)
            doc.save(output_path)

            if error_entries:
                report_path = self._write_error_report(report_path, self.task_snapshot, error_entries)

            self.progress_update.emit(100, "处理完成")
            self.finished.emit(len(error_entries) == 0, output_path, report_path or "")

        except Exception as e:
            report_path = self._write_error_report(
                report_path,
                self.task_snapshot,
                error_entries,
                traceback.format_exc(),
            )
            self.error.emit(str(e))
            self.error.emit(report_path or "")

    def _build_error_report_path(self, task_snapshot):
        output_path = task_snapshot.get("output_path") or task_snapshot.get("input_path")
        base_dir = os.path.dirname(output_path) or os.getcwd()
        base_name = os.path.splitext(os.path.basename(output_path))[0]
        return os.path.join(base_dir, f"{base_name}_error_report.txt")

    def _record_error(self, error_entries, stage, error, index=None, part_type=None, text=None):
        compact_text = re.sub(r"\s+", " ", text or "").strip()
        error_entries.append({
            "stage": stage,
            "index": index,
            "part_type": part_type,
            "text": compact_text[:200],
            "error": str(error),
        })

    def _write_error_report(self, report_path, task_snapshot, error_entries, fatal_trace=None):
        os.makedirs(os.path.dirname(report_path), exist_ok=True)
        with open(report_path, "w", encoding="utf-8") as report_file:
            report_file.write("论文格式助手错误报告\n")
            report_file.write("=" * 48 + "\n")
            report_file.write(f"输入文档: {task_snapshot.get('input_path', '')}\n")
            report_file.write(f"输出文档: {task_snapshot.get('output_path', '')}\n")
            report_file.write(f"问题数量: {len(error_entries)}\n\n")

            if error_entries:
                report_file.write("详细问题\n")
                report_file.write("-" * 48 + "\n")
                for idx, entry in enumerate(error_entries, 1):
                    report_file.write(f"{idx}. 阶段: {entry['stage']}\n")
                    if entry.get("index") is not None:
                        report_file.write(f"   序号: {entry['index']}\n")
                    if entry.get("part_type"):
                        report_file.write(f"   识别类型: {entry['part_type']}\n")
                    if entry.get("text"):
                        report_file.write(f"   文本片段: {entry['text']}\n")
                    report_file.write(f"   错误信息: {entry['error']}\n\n")

            if fatal_trace:
                report_file.write("致命错误堆栈\n")
                report_file.write("-" * 48 + "\n")
                report_file.write(fatal_trace)

        return report_path


class ThesisFormatterApp(QMainWindow):
    """论文格式助手主程序"""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("论文格式助手 v1.0")
        self.resize(1360, 860)
        self.setMinimumSize(1220, 780)

        # 设置应用样式
        self.setup_styles()

        # 存储格式规则
        self.format_vars = {}
        self.format_widgets = {}

        # 存储文件路径
        self.last_output_path = None
        self.last_report_path = None

        # 创建UI
        self.create_ui()

    def setup_styles(self):
        """设置应用样式 - 像素级复刻 Win11 现代卡片风"""
        QApplication.setStyle(QStyleFactory.create("Fusion"))

        stylesheet = """
            /* 1. 整体窗口背景：浅蓝灰色 */
            QMainWindow { background-color: #F0F4F8; }

            /* 2. 全局文字 */
            QWidget {
                color: #334155;
                font-family: "Microsoft YaHei UI", "Segoe UI", sans-serif;
                font-size: 13px;
            }

            /* 3. 两大块卡片背景（纯白+圆角+极淡的边框模拟阴影） */
            QFrame {
                background-color: #FFFFFF;
                border: 1px solid #E2E8F0;
                border-radius: 12px;
            }

            /* 4. 标题文字大小 */
            QLabel { border: none; background: transparent; }

            /* 5. 输入框样式 */
            QLineEdit {
                background-color: #FFFFFF;
                border: 1px solid #CBD5E1;
                border-radius: 6px;
                padding: 8px 12px;
                color: #334155;
            }
            QLineEdit:focus { border: 2px solid #60A5FA; }

            /* 6. 核心按钮：“开始排版”（亮蓝色渐变） */
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #3B82F6, stop:1 #2563EB);
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover { background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #60A5FA, stop:1 #3B82F6); }
            QPushButton:disabled { background: #CBD5E1; color: #94A3B8; }

            /* 7. 浏览按钮（浅蓝色） */
            QPushButton[objectName="browse"] {
                background-color: #93C5FD;
                color: white;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton[objectName="browse"]:hover { background-color: #60A5FA; }

            /* 8. 底部小按钮（浅灰色） */
            QPushButton[objectName="action"] {
                background-color: #E2E8F0;
                color: #475569;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton[objectName="action"]:hover { background-color: #CBD5E1; }
            QPushButton[objectName="action"]:disabled { background-color: #F1F5F9; color: #94A3B8; }

            /* 9. 表格整体样式（无网格线，斑马纹，直角） */
            QTableWidget {
                background-color: #FFFFFF;
                border: none;
                border-radius: 0px;
                gridline-color: transparent;
                alternate-background-color: #F8FAFC;
                selection-background-color: #EFF6FF;
                selection-color: #1E3A8A;
            }
            QTableWidget::item { border-bottom: 1px solid #F1F5F9; }

            /* 10. 表头样式（浅蓝灰底色） */
            QHeaderView::section {
                background-color: #E2EBF5;
                color: #475569;
                font-weight: bold;
                border: none;
                border-radius: 0px;
                padding: 8px;
            }

            /* 左上角交界处样式 */
            QTableWidget QTableCornerButton::section {
                background-color: #E2EBF5;
                border: none;
                border-radius: 0px;
            }

            /* 11. 进度条（胶囊形状） */
            QProgressBar {
                border: none;
                background-color: #E2E8F0;
                border-radius: 4px;
                height: 8px;
                text-align: center;
                color: transparent;
            }
            QProgressBar::chunk { background-color: #3B82F6; border-radius: 4px; }
            
            /* 12. 滚动条美化 */
            QScrollBar:vertical {
                border: none;
                background: #F8FAFC;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background: #CBD5E1;
                min-height: 20px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical:hover {
                background: #94A3B8;
            }

            QScrollBar:horizontal {
                border: none;
                background: #F8FAFC;
                height: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:horizontal {
                background: #CBD5E1;
                min-width: 20px;
                border-radius: 6px;
            }
            QScrollBar::handle:horizontal:hover {
                background: #94A3B8;
            }
        """
        self.setStyleSheet(stylesheet)

    def create_ui(self):
        """创建主界面 - Notion风格极简设计"""
        # 主窗口
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # 主布局
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(28, 28, 28, 28)

        # 标题区域
        self.create_header(main_layout)

        # 内容区域（左右布局）
        content_layout = QHBoxLayout()
        content_layout.setSpacing(20)

        # 左侧：格式规则表格
        format_card = self.create_card("格式规则设置")
        self.create_format_panel(format_card.layout())
        content_layout.addWidget(format_card, 3)

        # 右侧：文件操作面板
        file_card = self.create_card("文件操作")
        self.create_file_panel(file_card.layout())
        content_layout.addWidget(file_card, 1)

        main_layout.addLayout(content_layout)

    def create_header(self, parent_layout):
        """创建标题区域 - 清新学术风格"""
        header_widget = QWidget()
        header_layout = QVBoxLayout(header_widget)
        header_layout.setSpacing(8)
        header_layout.setContentsMargins(0, 0, 0, 0)

        # 主标题 - 使用明亮的蓝色和清晰的现代字体
        title_label = QLabel("论文格式助手")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 26px;
                font-weight: 700;
                color: #1e293b;
                letter-spacing: 2px;
            }
        """)

        header_layout.addWidget(title_label)
        parent_layout.addWidget(header_widget)
        parent_layout.addSpacing(16)

    def create_card(self, title):
        """创建卡片容器 - 清新学术风格"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)

        layout = QVBoxLayout(card)
        layout.setSpacing(16)
        layout.setContentsMargins(20, 20, 20, 20)

        # 卡片标题
        if title:
            title_label = QLabel(title)
            title_label.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #3b82f6;
                    padding-bottom: 4px;
                    background-color: transparent;
                    border: none;
                    letter-spacing: 0.5px;
                }
            """)
            layout.addWidget(title_label)

        return card

    def create_format_panel(self, parent_layout):
        """创建格式设置面板"""
        # 创建表格
        self.table = QTableWidget()
        self.table.setFocusPolicy(Qt.NoFocus)

        # 设置表格列数和行数（减1因为第一列移到垂直表头）
        self.table.setColumnCount(len(TABLE_COLUMNS) - 1)
        self.table.setRowCount(len(PARTS))

        # 设置水平表头（跳过第一列）
        headers = [col[1] for col in TABLE_COLUMNS[1:]]
        self.table.setHorizontalHeaderLabels(headers)

        # 固定表头
        self.table.horizontalHeader().setStretchLastSection(False)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)
        # 表头稍微调高一点
        self.table.horizontalHeader().setFixedHeight(45)

        # 设置垂直表头（显示论文部分名称）
        self.table.verticalHeader().setVisible(True)
        vertical_headers = [part_name for part_name, _ in PARTS]
        self.table.setVerticalHeaderLabels(vertical_headers)
        self.table.verticalHeader().setFixedWidth(150)  # 设置为原第一列的宽度
        self.table.verticalHeader().setDefaultAlignment(Qt.AlignLeft | Qt.AlignVCenter)

        # 设置列宽（跳过第一列）
        for col_idx, (_, _, width) in enumerate(TABLE_COLUMNS[1:]):
            self.table.setColumnWidth(col_idx, width)

        # 设置行高（从原来的40改成了52，避免下拉框太挤）
        self.table.verticalHeader().setDefaultSectionSize(52)

        # 填充表格内容
        self._fill_format_table()

        parent_layout.addWidget(self.table)

    def _fill_format_table(self):
        """填充格式表格"""
        default_formats = self._fill_format_default_formats()

        for row_idx, (part_name, part_key) in enumerate(PARTS):
            self.format_vars[part_key] = {}
            self.format_widgets[part_key] = {}

            defaults = default_formats.get(part_key, {})

            # 从第二列开始填充（第一列已经在垂直表头中）
            for col_idx, (field_name, _, _) in enumerate(TABLE_COLUMNS[1:]):
                widget = self._create_format_widget(part_key, field_name, defaults)
                if widget:
                    self.table.setCellWidget(row_idx, col_idx, widget)
                    self.format_widgets[part_key][field_name] = widget

    def _fill_format_default_formats(self):
        """获取默认格式设置"""
        return {
            "abstract_title": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "三号",
                              "alignment": "居中", "line_spacing": "1.5倍行距", "space_before": "0磅",
                              "space_after": "12磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "abstract_content": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "小四",
                                "alignment": "两端对齐", "line_spacing": "1.5倍行距", "space_before": "0磅",
                                "space_after": "0磅", "first_line_indent": "2字符", "overall_indent": "无", "bold": False, "italic": False},
            "heading1": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "三号",
                        "alignment": "居中", "line_spacing": "1.5倍行距", "space_before": "24磅",
                        "space_after": "12磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "heading2": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "四号",
                        "alignment": "左对齐", "line_spacing": "1.5倍行距", "space_before": "18磅",
                        "space_after": "6磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "heading3": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "小四",
                        "alignment": "左对齐", "line_spacing": "1.5倍行距", "space_before": "12磅",
                        "space_after": "6磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "heading4": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "小四",
                        "alignment": "左对齐", "line_spacing": "1.5倍行距", "space_before": "6磅",
                        "space_after": "6磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "body": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "小四",
                    "alignment": "两端对齐", "line_spacing": "1.5倍行距", "space_before": "0磅",
                    "space_after": "0磅", "first_line_indent": "2字符", "overall_indent": "无", "bold": False, "italic": False},
            "figure_caption": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "五号",
                              "alignment": "居中", "line_spacing": "单倍行距", "space_before": "6磅",
                              "space_after": "6磅", "first_line_indent": "无", "overall_indent": "无", "bold": False, "italic": False},
            "table_caption": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "五号",
                             "alignment": "居中", "line_spacing": "单倍行距", "space_before": "6磅",
                             "space_after": "6磅", "first_line_indent": "无", "overall_indent": "无", "bold": False, "italic": False},
            "table_content": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "五号",
                             "alignment": "居中", "line_spacing": "单倍行距", "space_before": "0磅",
                             "space_after": "0磅", "first_line_indent": "无", "overall_indent": "无", "bold": False, "italic": False},
            "table_note": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "小五",
                          "alignment": "左对齐", "line_spacing": "单倍行距", "space_before": "0磅",
                          "space_after": "0磅", "first_line_indent": "无", "overall_indent": "无", "bold": False, "italic": False},
            "formula": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "小四",
                       "alignment": "居中", "line_spacing": "1.5倍行距", "space_before": "6磅",
                       "space_after": "6磅", "first_line_indent": "无", "overall_indent": "无", "bold": False, "italic": False},
            "ref_title": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "三号",
                         "alignment": "居中", "line_spacing": "1.5倍行距", "space_before": "24磅",
                         "space_after": "12磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "ref_content": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "五号",
                           "alignment": "两端对齐", "line_spacing": "单倍行距", "space_before": "0磅",
                           "space_after": "0磅", "first_line_indent": "无", "overall_indent": "无", "bold": False, "italic": False},
            "ack_title": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "三号",
                         "alignment": "居中", "line_spacing": "1.5倍行距", "space_before": "24磅",
                         "space_after": "12磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "ack_content": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "小四",
                           "alignment": "两端对齐", "line_spacing": "1.5倍行距", "space_before": "0磅",
                           "space_after": "0磅", "first_line_indent": "2字符", "overall_indent": "无", "bold": False, "italic": False},
            "appendix_title": {"chinese_font": "黑体", "english_font": "Times New Roman", "font_size": "三号",
                              "alignment": "居中", "line_spacing": "1.5倍行距", "space_before": "24磅",
                              "space_after": "12磅", "first_line_indent": "无", "overall_indent": "无", "bold": True, "italic": False},
            "appendix_content": {"chinese_font": "宋体", "english_font": "Times New Roman", "font_size": "小四",
                                "alignment": "两端对齐", "line_spacing": "1.5倍行距", "space_before": "0磅",
                                "space_after": "0磅", "first_line_indent": "2字符", "overall_indent": "无", "bold": False, "italic": False},
        }


    def _create_format_widget(self, part_key, field_name, defaults):
        """创建格式设置组件"""
        widget = None
        default_value = defaults.get(field_name, "")

        if field_name == "chinese_font":
            widget = TableComboBox(CHINESE_FONTS, default_value)
            self.format_vars[part_key]["chinese_font"] = widget.combo
        elif field_name == "english_font":
            widget = TableComboBox(ENGLISH_FONTS, default_value)
            self.format_vars[part_key]["english_font"] = widget.combo
        elif field_name == "font_size":
            widget = TableComboBox(FONT_SIZE_NAMES, default_value)
            self.format_vars[part_key]["font_size"] = widget.combo
        elif field_name == "alignment":
            widget = TableComboBox(ALIGNMENT_NAMES, default_value)
            self.format_vars[part_key]["alignment"] = widget.combo
        elif field_name == "line_spacing":
            widget = TableComboBox(LINE_SPACING_NAMES, default_value)
            self.format_vars[part_key]["line_spacing"] = widget.combo
        elif field_name == "space_before":
            widget = TableComboBox(SPACINGS, default_value)
            self.format_vars[part_key]["space_before"] = widget.combo
        elif field_name == "space_after":
            widget = TableComboBox(SPACINGS, default_value)
            self.format_vars[part_key]["space_after"] = widget.combo
        elif field_name == "first_line_indent":
            widget = TableComboBox(FIRST_LINE_INDENTS, default_value)
            self.format_vars[part_key]["first_line_indent"] = widget.combo
        elif field_name == "overall_indent":
            widget = TableComboBox(OVERALL_INDENTS, default_value)
            self.format_vars[part_key]["overall_indent"] = widget.combo
        elif field_name == "bold":
            widget = TableCheckBox(default_value)
            self.format_vars[part_key]["bold"] = widget.checkbox
        elif field_name == "italic":
            widget = TableCheckBox(default_value)
            self.format_vars[part_key]["italic"] = widget.checkbox

        return widget

    def create_file_panel(self, parent_layout):
        """创建文件操作面板"""
        # 文件选择
        self._create_file_input(parent_layout, "选择论文文件", "input_file_entry", "input_file_btn")

        # 输出目录
        self._create_file_input(parent_layout, "输出文件位置", "output_dir_entry", "output_dir_btn", DEFAULT_OUTPUT_DIR_HINT)

        # 自定义文件名
        self._create_file_input(parent_layout, "自定义文件名", "filename_entry", None, DEFAULT_FILENAME_HINT)

        # 开始排版按钮
        self.start_btn = QPushButton("开始排版")
        self.start_btn.setMinimumHeight(42)
        self.start_btn.clicked.connect(self.start_formatting)
        parent_layout.addWidget(self.start_btn)

        parent_layout.addSpacing(20)

        # 输出操作按钮
        actions_layout = QHBoxLayout()
        actions_layout.setSpacing(10)

        self.open_doc_btn = QPushButton("打开文档")
        self.open_doc_btn.setEnabled(False)
        self.open_doc_btn.setObjectName("action")
        self.open_doc_btn.clicked.connect(self.open_output_document)
        actions_layout.addWidget(self.open_doc_btn)

        self.open_folder_btn = QPushButton("打开文件夹")
        self.open_folder_btn.setEnabled(False)
        self.open_folder_btn.setObjectName("action")
        self.open_folder_btn.clicked.connect(self.open_output_folder)
        actions_layout.addWidget(self.open_folder_btn)

        parent_layout.addLayout(actions_layout)

        parent_layout.addSpacing(20)

        # 进度区域
        self.progress = QProgressBar()
        self.progress.setTextVisible(False)
        self.progress.setMinimumHeight(4)
        parent_layout.addWidget(self.progress)

        self.progress_text = QLabel("准备就绪，请选择文件开始处理")
        self.progress_text.setStyleSheet("""
            QLabel {
                font-size: 10px;
                color: #9ca3af;
                background-color: transparent;
                border: none;
            }
        """)
        parent_layout.addWidget(self.progress_text)

    def _create_file_input(self, parent_layout, label_text, entry_name, btn_name, placeholder=""):
        """创建文件输入控件 - 清新学术风格"""
        label = QLabel(label_text)
        label.setStyleSheet("""
            QLabel {
                font-size: 12px;
                color: #475569;
                font-weight: 700;
                background-color: transparent;
                border: none;
                letter-spacing: 0.5px;
            }
        """)
        parent_layout.addWidget(label)

        input_layout = QHBoxLayout()
        input_layout.setSpacing(10)

        entry = QLineEdit()
        entry.setPlaceholderText(placeholder)
        setattr(self, entry_name, entry)
        input_layout.addWidget(entry)

        if btn_name:
            btn = QPushButton("浏览")
            btn.setMinimumWidth(70)
            btn.setMinimumHeight(36)
            btn.setObjectName("browse")
            if btn_name == "input_file_btn":
                btn.clicked.connect(self.select_file)
            elif btn_name == "output_dir_btn":
                btn.clicked.connect(self.select_output)
            setattr(self, btn_name, btn)
            input_layout.addWidget(btn)

        parent_layout.addLayout(input_layout)

    def select_file(self):
        """选择输入文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择论文文件",
            "",
            "Word文档 (*.docx);;所有文件 (*.*)"
        )
        if file_path:
            self.input_file_entry.setText(file_path)

    def select_output(self):
        """选择输出目录"""
        dir_path = QFileDialog.getExistingDirectory(self, "选择输出目录", "")
        if dir_path:
            self.output_dir_entry.setText(dir_path)

    def get_format_rules(self):
        """获取格式规则"""
        rules = {}
        for part_key, vars_dict in self.format_vars.items():
            rules[part_key] = {
                "chinese_font": vars_dict["chinese_font"].currentText(),
                "english_font": vars_dict["english_font"].currentText(),
                "font_size": vars_dict["font_size"].currentText(),
                "alignment": vars_dict["alignment"].currentText(),
                "line_spacing": vars_dict["line_spacing"].currentText(),
                "space_before": vars_dict["space_before"].currentText(),
                "space_after": vars_dict["space_after"].currentText(),
                "first_line_indent": vars_dict["first_line_indent"].currentText(),
                "overall_indent": vars_dict["overall_indent"].currentText(),
                "bold": vars_dict["bold"].isChecked(),
                "italic": vars_dict["italic"].isChecked(),
            }
        return rules

    def _normalize_output_dir(self, output_dir):
        if not output_dir or output_dir == DEFAULT_OUTPUT_DIR_HINT:
            return ""
        return output_dir

    def _normalize_custom_filename(self, custom_filename):
        if not custom_filename or custom_filename == DEFAULT_FILENAME_HINT:
            return ""
        return custom_filename

    def _build_output_path(self, input_path, output_dir, custom_filename):
        normalized_output_dir = self._normalize_output_dir(output_dir)
        normalized_filename = self._normalize_custom_filename(custom_filename)

        if normalized_output_dir:
            if normalized_filename:
                if not normalized_filename.endswith(".docx"):
                    normalized_filename += ".docx"
                return os.path.join(normalized_output_dir, normalized_filename)
            return os.path.join(normalized_output_dir, os.path.basename(input_path))

        name, ext = os.path.splitext(input_path)
        if normalized_filename:
            if not normalized_filename.endswith(".docx"):
                normalized_filename += ".docx"
            return os.path.join(os.path.dirname(input_path), normalized_filename)
        return f"{name}_formatted{ext}"

    def start_formatting(self):
        """开始排版"""
        input_path = self.input_file_entry.text().strip()
        if not input_path:
            QMessageBox.warning(self, "提示", "请先选择论文文件")
            return

        if not os.path.exists(input_path):
            QMessageBox.critical(self, "错误", "文件不存在")
            return

        output_dir = self.output_dir_entry.text().strip()
        normalized_output_dir = self._normalize_output_dir(output_dir)
        if normalized_output_dir and not os.path.isdir(normalized_output_dir):
            QMessageBox.critical(self, "错误", "输出文件夹不存在，请重新选择。")
            return

        custom_filename = self.filename_entry.text().strip()

        task_snapshot = {
            "input_path": input_path,
            "output_dir": output_dir,
            "custom_filename": custom_filename,
            "rules": self.get_format_rules(),
            "output_path": self._build_output_path(input_path, output_dir, custom_filename),
        }

        self.last_output_path = None
        self.last_report_path = None
        self.start_btn.setEnabled(False)
        self.open_doc_btn.setEnabled(False)
        self.open_folder_btn.setEnabled(False)
        self.progress.setValue(0)
        self.progress_text.setText("准备开始处理...")

        # 创建并启动工作线程
        self.format_thread = FormatThread(task_snapshot)
        self.format_thread.progress_update.connect(self._on_progress_update)
        self.format_thread.finished.connect(self._on_format_finished)
        self.format_thread.error.connect(self._on_format_error)
        self.format_thread.start()

    def _on_progress_update(self, value, text):
        """进度更新回调"""
        self.progress.setValue(value)
        self.progress_text.setText(text)

    def _on_format_finished(self, success, output_path, report_path):
        """格式完成回调"""
        self.last_output_path = output_path
        self.last_report_path = report_path

        if success:
            self.progress_text.setText("排版完成！")
        else:
            QMessageBox.warning(
                self,
                "处理完成",
                f"文档已生成，但处理过程中发现问题。\n错误报告已保存到：\n{report_path}",
            )
            self.progress_text.setText(f"处理完成，发现问题，已生成错误报告。")

        self.start_btn.setEnabled(True)
        self.open_doc_btn.setEnabled(True)
        self.open_folder_btn.setEnabled(True)

    def _on_format_error(self, error_msg):
        """格式错误回调"""
        QMessageBox.critical(
            self,
            "错误",
            f"处理失败: {error_msg}",
        )
        self.progress_text.setText("处理失败")
        self.start_btn.setEnabled(True)
        self.open_doc_btn.setEnabled(False)
        self.open_folder_btn.setEnabled(False)

    def open_output_document(self):
        """打开最后一次输出的文档"""
        if not self.last_output_path or not os.path.exists(self.last_output_path):
            QMessageBox.warning(self, "提示", "暂时没有可打开的输出文档")
            return

        try:
            os.startfile(self.last_output_path)
        except Exception as exc:
            QMessageBox.critical(self, "错误", f"打开文档失败: {exc}")

    def open_output_folder(self):
        """打开最后一次输出文件所在的文件夹"""
        if not self.last_output_path or not os.path.exists(self.last_output_path):
            QMessageBox.warning(self, "提示", "暂时没有可打开的输出文件夹")
            return

        try:
            subprocess.Popen(f'explorer /select,"{os.path.normpath(self.last_output_path)}"')
        except Exception:
            try:
                os.startfile(os.path.dirname(self.last_output_path))
            except Exception as exc:
                QMessageBox.critical(self, "错误", f"打开文件夹失败: {exc}")


# ==================== 主程序入口 ====================

def main():
    app = QApplication([])
    
    # 强制设置全局抗锯齿字体，让文字像图片里一样清晰
    font = app.font()
    font.setFamily("Microsoft YaHei UI") 
    font.setPointSize(10)
    app.setFont(font)

    window = ThesisFormatterApp()
    window.show()
    app.exec()

if __name__ == "__main__":
    main()
