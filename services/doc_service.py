# -*- coding: utf-8 -*-
"""
文档处理服务：解析 -> 识别 -> 格式应用 -> 保存
"""

import os
import re
import traceback
from typing import Callable, Dict, List, Optional

from docx import Document

from core.formatter import FormatApplier
from core.fusion_engine import create_fusion_identifier
from core.parser import DocumentParser


class DocumentProcessingService:
    """统一封装文档处理主流程"""

    def __init__(self, progress_callback: Optional[Callable[[int, str], None]] = None):
        self.progress_callback = progress_callback

    def process(self, task_snapshot: Dict) -> Dict:
        """
        处理文档并返回:
        {
            "success": bool,
            "output_path": str,
            "report_path": str,
            "error": str,
            "identify_result": dict
        }
        """
        error_entries: List[Dict] = []
        report_path = self._build_error_report_path(task_snapshot)

        try:
            input_path = task_snapshot["input_path"]
            output_path = task_snapshot["output_path"]
            rules = task_snapshot["rules"]
            ai_api_key = (task_snapshot.get("ai_api_key") or "").strip()
            ai_model = (task_snapshot.get("ai_model") or "").strip()

            self._emit_progress(10, "正在读取文档...")
            doc = Document(input_path)

            self._emit_progress(20, "正在识别文档结构...")
            parser = DocumentParser()
            units = parser.parse_document(doc)

            self._emit_progress(24, "规则识别中...")
            identifier = create_fusion_identifier(api_key=ai_api_key, model=ai_model)
            identify_result = identifier.identify_document(
                units,
                stage_callback=self._on_identify_stage,
            )
            final_labels = identify_result.get("final_labels", {})

            applier = FormatApplier(rules, doc)
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)

            if applier.style_manager:
                self._emit_progress(55, "正在建立样式...")
                for part_type, rule in rules.items():
                    try:
                        applier.style_manager.create_or_update_style(part_type, rule or {})
                    except Exception as exc:
                        self._record_error(
                            error_entries,
                            "样式创建",
                            exc,
                            part_type=part_type,
                        )
            else:
                self._emit_progress(55, "样式管理器不可用，使用回退格式...")

            self._emit_progress(60, "正在匹配段落样式...")
            for i, paragraph in enumerate(doc.paragraphs):
                try:
                    part_type = final_labels.get(i) or "body"
                    applier.apply_to_paragraph(paragraph, part_type)
                except Exception as exc:
                    self._record_error(
                        error_entries,
                        "段落样式应用",
                        exc,
                        index=i + 1,
                        part_type=final_labels.get(i) or "body",
                        text=paragraph.text,
                    )

                if total_paragraphs > 0 and (i % 100 == 0 or i == total_paragraphs - 1):
                    progress_val = 60 + int(((i + 1) / total_paragraphs) * 24)
                    self._emit_progress(
                        progress_val,
                        f"正在处理段落 {i + 1}/{total_paragraphs}...",
                    )

            if total_paragraphs == 0:
                self._emit_progress(84, "无段落可处理，已跳过。")

            self._emit_progress(85, "正在处理表格...")
            for i, table in enumerate(doc.tables):
                try:
                    applier.apply_to_table(table, rules)
                except Exception as exc:
                    table_preview = ""
                    try:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    table_preview = cell.text.strip()
                                    break
                            if table_preview:
                                break
                    except Exception:
                        table_preview = ""
                    self._record_error(
                        error_entries,
                        "表格处理",
                        exc,
                        index=i + 1,
                        part_type="table",
                        text=table_preview,
                    )

                progress = 85 + int((i + 1) / max(total_tables, 1) * 8)
                self._emit_progress(progress, f"正在处理表格 {i + 1}/{total_tables}...")

            self._emit_progress(93, "表格处理完成")
            if total_tables == 0:
                self._emit_progress(93, "无表格可处理，已跳过。")

            self._emit_progress(95, "正在保存文档...")
            os.makedirs(os.path.dirname(output_path) or os.getcwd(), exist_ok=True)
            doc.save(output_path)

            if error_entries:
                self._emit_progress(98, "正在写入错误报告...")
                report_path = self._write_error_report(
                    report_path,
                    task_snapshot,
                    error_entries,
                )
            else:
                report_path = ""

            self._emit_progress(100, "处理完成")
            return {
                "success": len(error_entries) == 0,
                "output_path": output_path,
                "report_path": report_path,
                "error": "",
                "identify_result": identify_result,
            }
        except Exception as exc:
            report_path = self._write_error_report(
                report_path,
                task_snapshot,
                error_entries,
                fatal_trace=traceback.format_exc(),
            )
            return {
                "success": False,
                "output_path": "",
                "report_path": report_path,
                "error": str(exc),
                "identify_result": {},
            }

    def _emit_progress(self, value: int, text: str):
        if self.progress_callback:
            self.progress_callback(value, text)

    def _on_identify_stage(self, stage: str, payload: Optional[Dict] = None):
        payload = payload or {}

        if stage == "rule_start":
            self._emit_progress(24, "规则识别中...")
            return
        if stage == "rule_done":
            self._emit_progress(30, "规则识别完成")
            return
        if stage == "anomaly_start":
            self._emit_progress(34, "异常检测中...")
            return
        if stage == "anomaly_done":
            if payload.get("need_ai"):
                reasons = payload.get("reasons") or []
                self._emit_progress(38, f"检测到异常，准备启动AI（原因数：{len(reasons)}）")
            else:
                self._emit_progress(38, "未检测到异常")
            return
        if stage == "ai_start":
            candidates = payload.get("candidates", 0)
            self._emit_progress(42, f"AI识别中（候选段落：{candidates}）...")
            return
        if stage == "ai_done":
            labeled = payload.get("labeled", 0)
            self._emit_progress(48, f"AI识别完成（已返回标签：{labeled}）")
            return
        if stage == "ai_skipped_disabled":
            self._emit_progress(46, "检测到异常，但未配置AI Key，保留规则结果")
            return
        if stage == "ai_skipped_no_candidates":
            self._emit_progress(46, "检测到异常，但无AI候选段落，保留规则结果")
            return
        if stage == "ai_not_needed":
            self._emit_progress(46, "无需AI识别，保留规则结果")
            return
        if stage == "identify_done":
            if payload.get("ai_used"):
                labeled = payload.get("ai_labeled_count", 0)
                self._emit_progress(50, f"文档结构识别完成（AI覆盖标签：{labeled}）")
            else:
                self._emit_progress(50, "文档结构识别完成（仅规则结果）")

    def _build_error_report_path(self, task_snapshot: Dict) -> str:
        output_path = task_snapshot.get("output_path") or task_snapshot.get("input_path")
        base_dir = os.path.dirname(output_path) or os.getcwd()
        base_name = os.path.splitext(os.path.basename(output_path))[0]
        return os.path.join(base_dir, f"{base_name}_error_report.txt")

    def _record_error(
        self,
        error_entries: List[Dict],
        stage: str,
        error: Exception,
        index: Optional[int] = None,
        part_type: Optional[str] = None,
        text: Optional[str] = None,
    ):
        compact_text = re.sub(r"\s+", " ", text or "").strip()
        error_entries.append(
            {
                "stage": stage,
                "index": index,
                "part_type": part_type,
                "text": compact_text[:200],
                "error": str(error),
            }
        )

    def _write_error_report(
        self,
        report_path: str,
        task_snapshot: Dict,
        error_entries: List[Dict],
        fatal_trace: Optional[str] = None,
    ) -> str:
        os.makedirs(os.path.dirname(report_path), exist_ok=True)
        with open(report_path, "w", encoding="utf-8") as report_file:
            report_file.write("论文格式助手错误报告\n")
            report_file.write("=" * 48 + "\n")
            report_file.write(f"输入文档: {task_snapshot.get('input_path', '')}\n")
            report_file.write(f"输出文档: {task_snapshot.get('output_path', '')}\n")
            report_file.write(f"问题数量: {len(error_entries)}\n\n")

            if error_entries:
                report_file.write("详细问题\n")
                report_file.write("-" * 48 + "\n")
                for idx, entry in enumerate(error_entries, 1):
                    report_file.write(f"{idx}. 阶段: {entry['stage']}\n")
                    if entry.get("index") is not None:
                        report_file.write(f"   序号: {entry['index']}\n")
                    if entry.get("part_type"):
                        report_file.write(f"   识别类型: {entry['part_type']}\n")
                    if entry.get("text"):
                        report_file.write(f"   文本片段: {entry['text']}\n")
                    report_file.write(f"   错误信息: {entry['error']}\n\n")

            if fatal_trace:
                report_file.write("致命错误堆栈\n")
                report_file.write("-" * 48 + "\n")
                report_file.write(fatal_trace)

        return report_path
